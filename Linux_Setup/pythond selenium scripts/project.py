import os
import urllib.request
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from docx import Document
from docx.shared import Inches
from selenium.webdriver.chrome.service import Service


binary_location = '/usr/bin/brave-browser'
driver_location = './chromedriver'
service = Service(driver_location)

# Set up the Chrome driver in headless mode
chrome_options = webdriver.ChromeOptions()
chrome_options.binary_location= binary_location
chrome_options.add_argument("--headless")
driver = webdriver.Chrome(service=service,options=chrome_options)

# Navigate to the BBC Sport website
driver.get("https://www.bbc.com/sport")

# Find the top 5 articles
try:
    articles = WebDriverWait(driver, 30).until(
        EC.presence_of_all_elements_located((By.CSS_SELECTOR, "div.top-story__content"))
    )[:5]
except TimeoutException:
    print("Timed out waiting for the top articles to load.")
    driver.quit()
    exit()
# Set up the document file
document = Document()
document.add_heading("Top 5 Articles on BBC Sport", 0)

# Loop through the articles and save the headline, link, and thumbnail to the document
for article in articles:
    headline = article.find_element(By.CSS_SELECTOR, "h3.top-story__title").text
    link = article.find_element(By.CSS_SELECTOR, "a").get_attribute("href")
    thumbnail_url = article.find_element(By.CSS_SELECTOR, "img").get_attribute("src")
    thumbnail_filename = os.path.basename(thumbnail_url)
    thumbnail_path = os.path.join(os.path.expanduser("~"), thumbnail_filename)
    
    # Download the thumbnail
    urllib.request.urlretrieve(thumbnail_url, thumbnail_path)
    
    # Add the article information to the document
    document.add_paragraph(headline, style="Heading 1")
    document.add_picture(thumbnail_path, width=Inches(1.25))
    document.add_paragraph(link)
    
    # Delete the thumbnail
    os.remove(thumbnail_path)

# Save the document file
document_filename = "top_5_articles.docx"
document_path = os.path.join(os.path.expanduser("~"), document_filename)
document.save(document_path)

# Close the browser
driver.quit()

print(f"The top 5 articles on BBC Sport have been saved to {document_path}.")
